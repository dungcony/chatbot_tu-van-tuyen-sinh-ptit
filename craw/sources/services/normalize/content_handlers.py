#!/usr/bin/env python3
"""
Xử lý các loại nội dung khi crawl:
- Dạng bảng (HTML table, ảnh bảng)
- Dạng ảnh
- Dạng video
- Dạng file (PDF, Word, Excel, ...)
"""
import io
import os
import re
from pathlib import Path
from urllib.parse import urlparse, urljoin

import requests

# ───────────────── Cấu hình ─────────────────
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
}

# Phân loại theo extension
FILE_EXTENSIONS = {
    ".pdf": "pdf",
    ".doc": "doc",
    ".docx": "docx",
    ".xls": "xls",
    ".xlsx": "xlsx",
    ".ppt": "ppt",
    ".pptx": "pptx",
}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}
VIDEO_EXTENSIONS = {".mp4", ".avi", ".mov", ".mkv", ".webm", ".m4v"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".ogg", ".m4a"}


def _get_content_type(url: str, timeout: int = 10) -> str | None:
    """Lấy Content-Type từ HEAD request."""
    try:
        r = requests.head(url, headers=HEADERS, timeout=timeout, allow_redirects=True)
        return (r.headers.get("Content-Type") or "").lower()
    except Exception:
        return None


def detect_content_type(url: str) -> str:
    """
    Phát hiện loại nội dung: html, table, image, video, file_pdf, file_doc, file_xlsx, ...
    Trả về: html | image | video | audio | pdf | docx | xlsx | unknown
    """
    parsed = urlparse(url)
    path = parsed.path.lower()
    ext = os.path.splitext(path)[1]

    if ext in IMAGE_EXTENSIONS:
        return "image"
    if ext in VIDEO_EXTENSIONS:
        return "video"
    if ext in AUDIO_EXTENSIONS:
        return "audio"
    if ext in FILE_EXTENSIONS:
        return FILE_EXTENSIONS[ext]

    # Không có extension rõ ràng -> có thể là HTML
    ct = _get_content_type(url)
    if ct:
        if "text/html" in ct:
            return "html"
        if "application/pdf" in ct:
            return "pdf"
        if "image/" in ct:
            return "image"
        if "video/" in ct:
            return "video"
        if "audio/" in ct:
            return "audio"
        if "application/vnd.openxmlformats" in ct or "application/msword" in ct:
            return "docx" if "word" in ct else "xlsx" if "sheet" in ct else "file"

    return "html"  # Mặc định coi là HTML


# ───────────────── Handler: PDF ─────────────────
def extract_text_pdf(url: str, timeout: int = 30) -> str | None:
    """Trích xuất text từ PDF. Cần: pip install pymupdf"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return None

    try:
        resp = requests.get(url, headers=HEADERS, timeout=timeout)
        resp.raise_for_status()
        doc = fitz.open(stream=resp.content, filetype="pdf")
        texts = []
        for page in doc:
            texts.append(page.get_text())
        doc.close()
        return "\n\n".join(texts).strip() if texts else None
    except Exception:
        return None


# ───────────────── Handler: Word DOCX ─────────────────
def extract_text_docx(url: str, timeout: int = 30) -> str | None:
    """Trích xuất text từ Word .docx. Cần: pip install python-docx"""
    try:
        from docx import Document
    except ImportError:
        return None

    try:
        resp = requests.get(url, headers=HEADERS, timeout=timeout)
        resp.raise_for_status()
        doc = Document(io.BytesIO(resp.content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts).strip() if parts else None
    except Exception:
        return None


# ───────────────── Handler: Excel XLSX ─────────────────
def extract_text_xlsx(url: str, timeout: int = 30) -> str | None:
    """Trích xuất text từ Excel .xlsx. Cần: pip install openpyxl"""
    try:
        import openpyxl
    except ImportError:
        return None

    try:
        resp = requests.get(url, headers=HEADERS, timeout=timeout)
        resp.raise_for_status()
        wb = openpyxl.load_workbook(io.BytesIO(resp.content), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                vals = [str(v).strip() if v is not None else "" for v in row]
                if any(vals):
                    rows.append(" | ".join(vals))
            if rows:
                parts.append(f"## {sheet.title}\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(parts).strip() if parts else None
    except Exception:
        return None


# ───────────────── Handler: Ảnh (OCR) ─────────────────
def extract_text_image(url: str, timeout: int = 30) -> str | None:
    """OCR ảnh lấy text. Cần: pip install easyocr torch. Dùng cho ảnh bảng điểm."""
    try:
        import sys
        _script_dir = Path(__file__).resolve().parent
        if str(_script_dir) not in sys.path:
            sys.path.insert(0, str(_script_dir))
        from ocr_image_table import ocr_table_from_image_url
    except ImportError:
        return None

    try:
        return ocr_table_from_image_url(url, verbose=False)
    except Exception:
        return None


# ───────────────── Handler: Video ─────────────────
def extract_metadata_video(url: str, page_html: str | None = None) -> dict:
    """
    Trích xuất metadata video: URL, embed, title.
    Nếu page_html có (trang chứa video), có thể parse thêm thông tin.
    """
    parsed = urlparse(url)
    return {
        "type": "video",
        "url": url,
        "domain": parsed.netloc,
        "path": parsed.path,
        "filename": os.path.basename(parsed.path),
    }


# ───────────────── Handler: Download file ─────────────────
def download_file(url: str, save_dir: Path, timeout: int = 60) -> Path | None:
    """Tải file về thư mục save_dir. Trả về đường dẫn file đã lưu."""
    try:
        resp = requests.get(url, headers=HEADERS, timeout=timeout, stream=True)
        resp.raise_for_status()
        parsed = urlparse(url)
        filename = os.path.basename(parsed.path) or "download"
        # Làm sạch tên file
        filename = re.sub(r"[^\w\-\.]", "_", filename)[:100]
        if not filename:
            filename = "download"
        save_dir.mkdir(parents=True, exist_ok=True)
        filepath = save_dir / filename
        with open(filepath, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                f.write(chunk)
        return filepath
    except Exception:
        return None


# ───────────────── API chính ─────────────────
def extract_content(url: str, content_type: str | None = None) -> tuple[str | None, str]:
    """
    Trích xuất nội dung text từ URL theo loại.
    Trả về (text, content_type).
    content_type: None = tự động phát hiện.
    """
    if content_type is None:
        content_type = detect_content_type(url)

    text = None
    if content_type == "pdf":
        text = extract_text_pdf(url)
    elif content_type in ("docx", "doc"):
        text = extract_text_docx(url)
    elif content_type in ("xlsx", "xls"):
        text = extract_text_xlsx(url)
    elif content_type == "image":
        text = extract_text_image(url)

    return text, content_type
